"""Helpers to build PDF/DOCX fixtures programmatically.

Used by tests/test_layout_parser.py — keeps binary blobs out of git and
ensures fixtures are reproducible.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


def make_simple_pdf(path: Path) -> Path:
    """Plain prose PDF, no headings, no tables."""
    doc = SimpleDocTemplate(str(path), pagesize=A4)
    styles = getSampleStyleSheet()
    story = [
        Paragraph(
            "Ini adalah paragraf biasa tanpa heading apapun. "
            "Cukup untuk uji parser tidak crash di dokumen sederhana.",
            styles["BodyText"],
        ),
        Spacer(1, 0.4 * cm),
        Paragraph(
            "Paragraf kedua melanjutkan tulisan biasa, juga tanpa heading.",
            styles["BodyText"],
        ),
    ]
    doc.build(story)
    return path


def make_pdf_with_table(path: Path) -> Path:
    """PDF with a small table inside."""
    doc = SimpleDocTemplate(str(path), pagesize=A4)
    styles = getSampleStyleSheet()
    story = [
        Paragraph("Ringkasan Anggaran", styles["BodyText"]),
        Spacer(1, 0.4 * cm),
        Table(
            [
                ["Item", "Jumlah", "Keterangan"],
                ["Gaji", "100jt", "Q1"],
                ["Tunjangan", "25jt", "Q1"],
                ["Operasional", "50jt", "Q2"],
            ],
            style=TableStyle(
                [
                    ("GRID", (0, 0), (-1, -1), 0.5, "#000000"),
                ]
            ),
        ),
        Spacer(1, 0.4 * cm),
        Paragraph("Catatan akhir setelah tabel.", styles["BodyText"]),
    ]
    doc.build(story)
    return path


def make_pdf_with_indonesian_heading(path: Path) -> Path:
    """PDF with formal Indonesian section structure (BAB / Pasal)."""
    doc = SimpleDocTemplate(str(path), pagesize=A4)
    styles = getSampleStyleSheet()
    heading = ParagraphStyle(
        "BabHeading",
        parent=styles["Heading1"],
        fontSize=18,
        spaceAfter=12,
    )
    story = [
        Paragraph("BAB I PENDAHULUAN", heading),
        Paragraph(
            "Latar belakang penyusunan dokumen ini bermula dari kebutuhan "
            "untuk mengatur kebijakan baru.",
            styles["BodyText"],
        ),
        Spacer(1, 0.4 * cm),
        Paragraph("Pasal 1", heading),
        Paragraph(
            "Dalam peraturan ini yang dimaksud dengan istilah berikut...",
            styles["BodyText"],
        ),
    ]
    doc.build(story)
    return path


def make_docx_with_heading_and_table(path: Path) -> Path:
    """DOCX with Word native Heading 1 / Heading 2 styles plus a table."""
    docx = DocxDocument()
    docx.add_heading("Bab 1: Pendahuluan", level=1)
    docx.add_paragraph("Paragraf pertama di bawah Bab 1.")
    docx.add_heading("1.1 Latar Belakang", level=2)
    docx.add_paragraph("Latar belakang dijelaskan di sini.")

    table = docx.add_table(rows=3, cols=3)
    header_cells = table.rows[0].cells
    header_cells[0].text = "Item"
    header_cells[1].text = "Jumlah"
    header_cells[2].text = "Keterangan"
    table.rows[1].cells[0].text = "Gaji"
    table.rows[1].cells[1].text = "100jt"
    table.rows[1].cells[2].text = "Q1"
    table.rows[2].cells[0].text = "Tunjangan"
    table.rows[2].cells[1].text = "25jt"
    table.rows[2].cells[2].text = "Q1"

    docx.add_heading("1.2 Tujuan", level=2)
    docx.add_paragraph("Tujuan dari kebijakan ini.")

    docx.save(str(path))
    return path


def make_corrupt_pdf(path: Path) -> Path:
    """Write a file that is NOT a valid PDF, to exercise fallback path."""
    path.write_bytes(b"this is not a real pdf at all")
    return path
