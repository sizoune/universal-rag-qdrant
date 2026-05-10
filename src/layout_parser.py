"""Layout-aware document parser for PDF and DOCX.

Produces a flat list of `Element` objects (paragraph / heading / table /
list_item), then `chunk_elements` groups them into LangChain `Document`
chunks while:

- Maintaining a heading path stack (level-aware push/pop)
- Prepending the current heading hierarchy to chunk content for retrieval
- Emitting tables as their own chunks (split per row when oversized,
  with the table header reprinted on every chunk)
- Preserving page metadata for PDF chunks

Heading detection on PDFs is hybrid: Indonesian formal-document regex
patterns (BAB / Pasal / 1.1) take precedence; font-size ratio relative
to the body font is the fallback.
"""

from __future__ import annotations

import logging
import re
from collections import Counter
from dataclasses import dataclass
from re import Pattern
from typing import Literal

import pdfplumber
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from langchain_core.documents import Document

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Public Element type
# ---------------------------------------------------------------------------

ElementKind = Literal["heading", "paragraph", "table", "list_item"]


@dataclass(frozen=True)
class Element:
    """A single structured fragment from a document.

    For headings, `level` is 1-6 (1 is most senior).
    For tables, `text` is a pre-built Markdown table string and
    `table_caption` is the optional caption.
    """

    kind: ElementKind
    level: int
    text: str
    page: int | None = None
    table_caption: str | None = None


# ---------------------------------------------------------------------------
# Indonesian heading patterns (regex-based fast path)
# ---------------------------------------------------------------------------

INDONESIAN_HEADING_PATTERNS: list[tuple[Pattern[str], int]] = [
    # Level 1 — chapter / annex / major part
    (re.compile(r"^BAB\s+[IVXLCDM]+\b", re.IGNORECASE), 1),
    (
        re.compile(
            r"^BAGIAN\s+(?:KESATU|KEDUA|KETIGA|KEEMPAT|KELIMA|KEENAM|KETUJUH)\b",
            re.IGNORECASE,
        ),
        1,
    ),
    (re.compile(r"^Bab\s+\d+\b"), 1),
    (re.compile(r"^Bagian\s+\d+\b"), 1),
    (re.compile(r"^Lampiran\b", re.IGNORECASE), 1),
    # Level 2 — articles / sections
    (re.compile(r"^Pasal\s+\d+\b"), 2),
    (re.compile(r"^[IVXLCDM]+\.\s+[A-Z]"), 2),
    (re.compile(r"^\d+\.\d+\s"), 2),
    # Level 3 — sub-sections
    (re.compile(r"^\d+\.\d+\.\d+\s"), 3),
    (re.compile(r"^[A-Z]\.\s+[A-Z]"), 3),
]


def detect_heading(
    text: str,
    font_size: float | None = None,
    body_font_size: float | None = None,
) -> tuple[bool, int]:
    """Return (is_heading, level).

    Tries Indonesian regex patterns first; falls back to a font-size ratio
    heuristic against `body_font_size` if both font sizes are provided.
    """
    if not text or not text.strip():
        return False, 0
    stripped = text.strip()

    for pattern, level in INDONESIAN_HEADING_PATTERNS:
        if pattern.match(stripped):
            return True, level

    if font_size is not None and body_font_size and body_font_size > 0:
        ratio = font_size / body_font_size
        if ratio >= 1.5:
            return True, 1
        if ratio >= 1.25:
            return True, 2
        if ratio >= 1.1:
            return True, 3

    return False, 0


# ---------------------------------------------------------------------------
# Chunker — pure function over a list of Elements
# ---------------------------------------------------------------------------


def chunk_elements(
    elements: list[Element],
    max_chunk_size: int = 1000,
) -> list[Document]:
    """Group elements into LangChain Documents respecting layout boundaries."""
    chunks: list[Document] = []
    heading_stack: list[tuple[int, str]] = []  # (level, text)
    buffer_parts: list[str] = []
    buffer_page: int | None = None
    buffer_size = 0

    def heading_path() -> list[str]:
        return [text for _, text in heading_stack]

    def flush():
        nonlocal buffer_parts, buffer_page, buffer_size
        if not buffer_parts:
            return
        body = "\n\n".join(buffer_parts)
        chunks.append(
            _build_chunk(
                body=body,
                heading_path=heading_path(),
                page=buffer_page,
                kind="paragraph",
            )
        )
        buffer_parts = []
        buffer_page = None
        buffer_size = 0

    for el in elements:
        if el.kind == "heading":
            flush()
            # Pop any heading at deeper-or-equal level than this one
            while heading_stack and heading_stack[-1][0] >= el.level:
                heading_stack.pop()
            heading_stack.append((el.level, el.text))
            continue

        if el.kind == "table":
            flush()
            for body in _split_table_body(el.text, el.table_caption, max_chunk_size):
                chunks.append(
                    _build_chunk(
                        body=body,
                        heading_path=heading_path(),
                        page=el.page,
                        kind="table",
                        table_caption=el.table_caption,
                    )
                )
            continue

        # Paragraph or list item — buffer up to max_chunk_size
        text = el.text.strip()
        if not text:
            continue
        if buffer_parts and buffer_size + len(text) > max_chunk_size:
            flush()
        buffer_parts.append(text)
        buffer_size += len(text)
        if buffer_page is None:
            buffer_page = el.page

    flush()
    return chunks


def _build_chunk(
    body: str,
    heading_path: list[str],
    page: int | None,
    kind: str,
    table_caption: str | None = None,
) -> Document:
    if heading_path:
        prefix = "# " + " / ".join(heading_path) + "\n\n"
        content = prefix + body
    else:
        content = body

    metadata: dict = {
        "heading_path": list(heading_path),
        "chunk_kind": kind,
        "parser_version": 2,
    }
    if page is not None:
        metadata["page"] = page
    if table_caption:
        metadata["table_caption"] = table_caption
    return Document(page_content=content, metadata=metadata)


def _split_table_body(
    table_md: str, caption: str | None, max_size: int
) -> list[str]:
    """Split a markdown table into chunks of at most `max_size` chars.

    Keeps the header row reprinted on every chunk so each chunk is
    self-contained for embedding/retrieval.
    """
    if len(table_md) <= max_size:
        return [_with_caption(table_md, caption, part=None, total=None)]

    lines = table_md.split("\n")
    if len(lines) < 3:
        # Not a recognizable markdown table; emit as a single chunk anyway.
        return [_with_caption(table_md, caption, part=None, total=None)]

    header, sep, *rows = lines
    fixed_size = len(header) + len(sep) + 2  # two newlines

    parts: list[list[str]] = []
    current: list[str] = []
    current_size = fixed_size

    for row in rows:
        if current and current_size + len(row) + 1 > max_size:
            parts.append(current)
            current = []
            current_size = fixed_size
        current.append(row)
        current_size += len(row) + 1

    if current:
        parts.append(current)

    total = len(parts)
    out = []
    for idx, rows_subset in enumerate(parts, start=1):
        body = "\n".join([header, sep] + rows_subset)
        out.append(
            _with_caption(
                body, caption, part=idx if total > 1 else None, total=total if total > 1 else None
            )
        )
    return out


def _with_caption(
    body: str, caption: str | None, part: int | None, total: int | None
) -> str:
    if not caption and part is None:
        return body
    if caption and part:
        return f"## Tabel: {caption} (chunk {part}/{total})\n\n{body}"
    if caption:
        return f"## Tabel: {caption}\n\n{body}"
    if part:
        return f"## Tabel (chunk {part}/{total})\n\n{body}"
    return body


# ---------------------------------------------------------------------------
# PDF parser
# ---------------------------------------------------------------------------


def parse_pdf(filepath: str) -> list[Element]:
    """Parse PDF into a flat list of Elements (paragraphs, headings, tables).

    Heading detection is hybrid: Indonesian regex first, then font-size
    ratio against the document's dominant body font size.
    """
    elements: list[Element] = []
    with pdfplumber.open(filepath) as pdf:
        body_font_size = _detect_body_font_size(pdf)

        for page_idx, page in enumerate(pdf.pages, start=1):
            tables = page.extract_tables() or []
            text = page.extract_text() or ""

            line_font_sizes = _line_font_sizes(page)

            for paragraph in _split_paragraphs(text):
                first_line = paragraph.split("\n", 1)[0].strip()
                font_size = line_font_sizes.get(first_line)
                is_heading, level = detect_heading(
                    first_line, font_size=font_size, body_font_size=body_font_size
                )

                if is_heading:
                    elements.append(
                        Element(
                            kind="heading",
                            level=level,
                            text=first_line,
                            page=page_idx,
                        )
                    )
                    rest = (
                        paragraph.split("\n", 1)[1].strip()
                        if "\n" in paragraph
                        else ""
                    )
                    if rest:
                        elements.append(
                            Element(
                                kind="paragraph",
                                level=0,
                                text=rest,
                                page=page_idx,
                            )
                        )
                else:
                    elements.append(
                        Element(
                            kind="paragraph",
                            level=0,
                            text=paragraph.strip(),
                            page=page_idx,
                        )
                    )

            for tbl in tables:
                if not tbl:
                    continue
                md = _table_to_markdown(tbl)
                if md:
                    elements.append(
                        Element(
                            kind="table",
                            level=0,
                            text=md,
                            page=page_idx,
                            table_caption=None,
                        )
                    )

    return elements


def _split_paragraphs(text: str) -> list[str]:
    """Split a page's text into paragraphs separated by blank lines."""
    if not text:
        return []
    blocks = re.split(r"\n\s*\n", text)
    return [b.strip() for b in blocks if b.strip()]


def _detect_body_font_size(pdf) -> float | None:
    """Find the most common (mode) font size across the whole document.

    Uses pdfplumber's `chars` attribute. Returns None if nothing detectable.
    """
    sizes: Counter[float] = Counter()
    for page in pdf.pages:
        for char in page.chars:
            size = char.get("size")
            if size:
                sizes[round(float(size), 1)] += 1
    if not sizes:
        return None
    return sizes.most_common(1)[0][0]


def _line_font_sizes(page) -> dict[str, float]:
    """Approximate font size per text line on a page (key: line text).

    pdfplumber doesn't directly expose lines + sizes together; we group
    chars by their `top` coordinate (rounded) and compute the mean size
    per group, then build text -> size mapping.
    """
    line_buckets: dict[float, list[dict]] = {}
    for char in page.chars:
        top = round(float(char.get("top", 0)), 0)
        line_buckets.setdefault(top, []).append(char)

    line_to_size: dict[str, float] = {}
    for chars_in_line in line_buckets.values():
        if not chars_in_line:
            continue
        text = "".join(c.get("text", "") for c in chars_in_line).strip()
        if not text:
            continue
        sizes = [float(c.get("size", 0)) for c in chars_in_line if c.get("size")]
        if not sizes:
            continue
        line_to_size[text] = sum(sizes) / len(sizes)
    return line_to_size


def _table_to_markdown(table: list[list[str | None]]) -> str:
    """Convert pdfplumber/python-docx table (rows of cells) to Markdown."""
    if not table:
        return ""
    cleaned = [[_clean_cell(c) for c in row] for row in table]
    # Drop fully-empty trailing/leading rows
    cleaned = [row for row in cleaned if any(cell for cell in row)]
    if not cleaned:
        return ""

    width = max(len(row) for row in cleaned)
    cleaned = [row + [""] * (width - len(row)) for row in cleaned]

    header = cleaned[0]
    body = cleaned[1:]
    md_header = "| " + " | ".join(header) + " |"
    md_sep = "|" + "|".join(["---"] * width) + "|"
    md_rows = ["| " + " | ".join(row) + " |" for row in body]
    return "\n".join([md_header, md_sep] + md_rows)


def _clean_cell(cell: str | None) -> str:
    if cell is None:
        return ""
    return cell.replace("|", "\\|").replace("\n", " ").strip()


# ---------------------------------------------------------------------------
# DOCX parser
# ---------------------------------------------------------------------------


def parse_docx(filepath: str) -> list[Element]:
    """Parse DOCX in document order: paragraphs (with style-based headings)
    and tables, interleaved as they appear in the body.
    """
    docx = DocxDocument(filepath)
    elements: list[Element] = []

    paragraphs_by_element = {p._element: p for p in docx.paragraphs}
    tables_by_element = {t._element: t for t in docx.tables}

    for child in docx.element.body.iterchildren():
        if child.tag == qn("w:p"):
            para = paragraphs_by_element.get(child)
            if para is None:
                continue
            text = (para.text or "").strip()
            if not text:
                continue
            level = _heading_level_from_docx_style(para.style.name or "")
            if level > 0:
                elements.append(
                    Element(kind="heading", level=level, text=text, page=None)
                )
            else:
                elements.append(
                    Element(kind="paragraph", level=0, text=text, page=None)
                )
        elif child.tag == qn("w:tbl"):
            tbl = tables_by_element.get(child)
            if tbl is None:
                continue
            rows = [
                [cell.text.strip() for cell in row.cells] for row in tbl.rows
            ]
            md = _table_to_markdown(rows)
            if md:
                elements.append(
                    Element(
                        kind="table",
                        level=0,
                        text=md,
                        page=None,
                        table_caption=None,
                    )
                )

    return elements


def _heading_level_from_docx_style(style_name: str) -> int:
    """Map Word style 'Heading N' to int level 1-9; 0 if not a heading style."""
    if not style_name:
        return 0
    name = style_name.strip()
    if not name.lower().startswith("heading"):
        return 0
    parts = name.split()
    if len(parts) >= 2 and parts[1].isdigit():
        level = int(parts[1])
        return level if 1 <= level <= 9 else 0
    if len(parts) == 1:
        return 1
    return 0
